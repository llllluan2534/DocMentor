from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging
from typing import List, Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)

class ExportService:
    """Service for exporting documents to various formats."""
    
    def __init__(self):
        self.export_dir = "exports"
        os.makedirs(self.export_dir, exist_ok=True)
        
    def export_summary_to_pdf(
        self,
        document_title: str,
        summary: str,
        metadata: Dict[str, Any] = None
    ) -> str:
        """
        Export summary to PDF
        
        Returns:
            Path to generated PDF file.
        """
        try:
            # Tạo tên file với timestamp để tránh trùng lặp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"summary_{timestamp}.pdf"
            filepath = os.path.join(self.export_dir, filename)
            
            # Tạo tài liệu PDF
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#4B8BBE'),
                spaceAfter=30,
                alignment=1  # Centered
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                leading=16,
                spaceAfter=20
            )
            
            # Thêm tiêu đề
            story.append(Paragraph("TÓM TẮT TÀI LIỆU", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Thêm nội dung tài liệu
            story.append(Paragraph(f"<b>Tài liệu:</b> {document_title}", body_style))
            
            if metadata:
                story.append(Paragraph(
                    f"<b>Độ dài:</b> {metadata.get('length', 'medium').title()}", 
                    body_style
                ))
                story.append(Paragraph(
                    f"<b>Số từ:</b> {metadata.get('word_count', 0)}", 
                    body_style
                ))
                
            story.append(Paragraph(
                f"<b>Ngày tạo:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", 
                body_style
            ))
            story.append(Spacer(1, 0.3*inch))
            
            # Thêm nội dung tóm tắt
            story.append(Paragraph("<b> NỘI DUNG TÓM TẮT:</b>", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            # Chia tóm tắt thành các đoạn nhỏ
            for para in summary.split('\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))
                    
            # Footer
            story.append(Spacer(1, 0.5*inch))
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.grey,
                alignment=1  # Centered
            )
            story.append(Paragraph(
                "Được tạo bởi DocMentor - AI Document Analysis System", 
                footer_style
            ))
            
            # Xây dựng tài liệu PDF
            doc.build(story)
            logger.info(f"PDF summary exported successfully: {filepath}")
            
            return filepath
        
        except Exception as e:
            logger.error(f"Error exporting PDF summary: {e}")
            raise
        
    def export_summary_to_docx(
        self,
        document_title: str,
        summary: str,
        metadata: Dict[str, Any] = None
    ) -> str:
        """
        Export summary to DOCX
        
        Returns:
            Path to generated DOCX file.
        """
        try:
            # Tạo tên file với timestamp để tránh trùng lặp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"summary_{timestamp}.docx"
            filepath = os.path.join(self.export_dir, filename)
            
            # Tạo tài liệu DOCX
            doc = DocxDocument()
            
            # Thêm tiêu đề
            title = doc.add_heading("TÓM TẮT TÀI LIỆU", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm nội dung tài liệu
            doc.add_paragraph()
            info_para = doc.add_paragraph()
            info_para.add_run("Tài liệu: ").bold = True
            info_para.add_run(document_title)
            
            if metadata:
                length_para = doc.add_paragraph()
                length_para.add_run("Độ dài: ").bold = True
                length_para.add_run(metadata.get('length', 'medium').title())
                
                count_para = doc.add_paragraph()
                count_para.add_run("Số từ: ").bold = True
                count_para.add_run(str(metadata.get('word_count', 0)))
                
            date_para = doc.add_paragraph()
            date_para.add_run("Ngày tạo: ").bold = True
            date_para.add_run(datetime.now().strftime('%d/%m/%Y %H:%M'))
            
            # Summary heading
            doc.add_paragraph()
            doc.add_heading("NỘI DUNG TÓM TẮT:", level=1)
            
            # Summary content
            for para in summary.split('\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.line_spacing = 1.5
                    
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph(
                "Được tạo bởi DocMentor - AI Document Analysis System"
            )
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer.runs[0]
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Lưu tài liệu DOCX
            doc.save(filepath)
            logger.info(f"DOCX summary exported successfully: {filepath}")
            
            return filepath
        
        except Exception as e:
            logger.error(f"Error exporting DOCX summary: {e}")
            raise
        
    def export_quiz_to_pdf(
        self,
        document_title: str,
        questions: List[Dict[str, Any]],
        difficulty: str,
        include_answers: bool = False
    ) -> str:
        """
        Export quiz to PDF
        
        Args:
            document_title: Title of document
            questions: List of quiz questions
            difficulty: Quiz difficulty level
            include_answers: Whether to include answer key
        
        Returns:
            Path to generated PDF file.
        """
        try:
            # Tạo tên file với timestamp để tránh trùng lặp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"quiz_{timestamp}.pdf"
            filepath = os.path.join(self.export_dir, filename)
            
            # Tạo tài liệu PDF
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'QuizTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#4B8BBE'),
                spaceAfter=20,
                alignment=1  # Centered
            )
            story.append(Paragraph("BÀI KIỂM TRA", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Info
            info_style = styles['Normal']
            story.append(Paragraph(f"<b>Tài liệu:</b> {document_title}", info_style))
            story.append(Paragraph(f"<b>Độ khó:</b> {difficulty.title()}", info_style))
            story.append(Paragraph(f"<b>Số câu hỏi:</b> {len(questions)}", info_style))
            story.append(Paragraph(
                f"<b>Ngày tạo:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", 
                info_style
            ))
            story.append(Spacer(1, 0.3*inch))
            
            # Questions
            question_style = ParagraphStyle(
                'Question',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=8,
                fontName='Helvetica-Bold'
            )
            
            option_style = ParagraphStyle(
                'Option',
                parent=styles['Normal'],
                fontSize=10,
                leftIndent=20,
                spaceAfter=4
            )
            
            for idx, q in enumerate(questions, 1):
                # Số câu hỏi và văn bản câu hỏi
                story.append(Paragraph(f"Câu {idx}: {q['question']}", question_style))
                
                # Các lựa chọn
                for option in q['options']:
                    story.append(Paragraph(option, option_style))
                    
                story.append(Spacer(1, 0.15*inch))
                
            # Answer Key (nếu có)
            if include_answers:
                story.append(PageBreak())
                story.append(Paragraph("ĐÁP ÁN", title_style))
                story.append(Spacer(1, 0.2*inch))
                
                for idx, q in enumerate(questions, 1):
                    answer_text = f"Câu {idx}: <b>{q['correct']}</b> - {q['explanation']}"
                    story.append(Paragraph(answer_text, styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                    
            # Xây dựng tài liệu PDF
            doc.build(story)
            logger.info(f"PDF quiz exported successfully: {filepath}")
            
            return filepath
        
        except Exception as e:
            logger.error(f"Error exporting PDF quiz: {e}")
            raise